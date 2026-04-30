import json
import html
import io
import re
from datetime import UTC, datetime
from pathlib import Path
from typing import Any
from urllib.parse import urljoin, urlparse

import httpx
import yaml

from app.testcase_gen.services.openapi_service import openapi_service

try:
    from bs4 import BeautifulSoup
except ImportError:  # pragma: no cover
    BeautifulSoup = None

OPENAPI_EXTENSIONS = {".json", ".yaml", ".yml"}
TEXT_EXTENSIONS = {".md", ".txt", ".csv", ".html", ".htm", ".docx", ".xlsx", ".xls"}
SUPPORTED_EXTENSIONS = OPENAPI_EXTENSIONS | TEXT_EXTENSIONS | {".har"}
HTTP_METHODS = {"GET", "POST", "PUT", "PATCH", "DELETE", "OPTIONS", "HEAD"}
DOC_DISCOVERY_PATHS = ["/openapi.json", "/swagger.json", "/v3/api-docs"]


def parse_api_description_file(file_path: str, *, filename: str | None = None) -> dict[str, Any]:
    path = Path(file_path)
    suffix = path.suffix.lower()
    name = filename or path.name

    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError("仅支持 OpenAPI / Postman / HAR / Markdown / Word / Excel / HTML / TXT / CSV 文件")

    content = path.read_bytes()
    api_info = _parse_api_description_content(content, suffix=suffix, source_name=name)
    return {
        "api_info": api_info,
        "parsed_at": datetime.now(UTC).isoformat().replace("+00:00", "Z"),
        "file_path": file_path,
    }


async def parse_api_description_url(
    url: str,
    *,
    client: httpx.AsyncClient,
    response: httpx.Response | None = None,
) -> dict[str, Any]:
    initial = response or await client.get(url)
    initial.raise_for_status()
    content_type = (initial.headers.get("content-type") or "").lower()
    suffix = Path(urlparse(str(initial.url)).path).suffix.lower()

    openapi_info = _try_parse_openapi(initial.content, suffix=suffix, source_name=url)
    if openapi_info:
        return openapi_info

    generic_json_info = _try_parse_json_formats(initial.content, source_name=url)
    if generic_json_info:
        return generic_json_info

    if "html" in content_type or suffix in {".html", ".htm", ""}:
        discovered = await _discover_openapi_from_html(str(initial.url), initial.text, client)
        if discovered:
            return discovered
        text = _html_to_text(initial.content)
        return _build_api_info_from_text(text, title=f"Imported from {url}", server_url=_server_from_url(str(initial.url)))

    text = initial.text
    return _build_api_info_from_text(text, title=f"Imported from {url}", server_url=_server_from_url(str(initial.url)))


def _parse_api_description_content(content: bytes, *, suffix: str, source_name: str) -> dict[str, Any]:
    openapi_info = _try_parse_openapi(content, suffix=suffix, source_name=source_name)
    if openapi_info:
        return openapi_info

    if suffix in {".json", ".har"}:
        json_info = _try_parse_json_formats(content, source_name=source_name)
        if json_info:
            return json_info

    if suffix in {".html", ".htm"}:
        text = _html_to_text(content)
        return _build_api_info_from_text(text, title=f"Imported from {source_name}")

    text = _extract_text(content, suffix=suffix, source_name=source_name)
    return _build_api_info_from_text(text, title=f"Imported from {source_name}")


def _try_parse_openapi(content: bytes, *, suffix: str, source_name: str) -> dict[str, Any] | None:
    if suffix not in OPENAPI_EXTENSIONS and suffix not in {".json", ""}:
        return None

    try:
        raw = content.decode("utf-8", errors="ignore")
        if suffix == ".json":
            data = json.loads(raw)
        elif suffix in {".yaml", ".yml"}:
            data = yaml.safe_load(raw)
        else:
            try:
                data = json.loads(raw)
            except Exception:
                data = yaml.safe_load(raw)
    except Exception:
        return None

    if not isinstance(data, dict) or "paths" not in data:
        return None
    if "openapi" not in data and "swagger" not in data:
        return None

    parsed = openapi_service._parse_api_specification(data)
    return _ensure_api_info_shape(parsed)


def _try_parse_json_formats(content: bytes, *, source_name: str) -> dict[str, Any] | None:
    try:
        data = json.loads(content.decode("utf-8", errors="ignore"))
    except Exception:
        return None

    if not isinstance(data, dict):
        return None

    postman_info = _parse_postman_collection(data, source_name=source_name)
    if postman_info:
        return postman_info

    har_info = _parse_har(data, source_name=source_name)
    if har_info:
        return har_info

    generic_info = _parse_generic_request_tree(data, source_name=source_name)
    if generic_info:
        return generic_info

    return None


def _parse_postman_collection(data: dict[str, Any], *, source_name: str) -> dict[str, Any] | None:
    schema = str(((data.get("info") or {}).get("schema")) or "")
    if "postman" not in schema.lower() and "item" not in data:
        return None

    operations: list[dict[str, Any]] = []
    for item in data.get("item", []):
        _walk_postman_items(item, operations, group_tags=[])

    if not operations:
        return None

    title = ((data.get("info") or {}).get("name")) or f"Postman Collection {source_name}"
    return _build_api_info_from_operations(operations, title=title)


def _walk_postman_items(item: dict[str, Any], operations: list[dict[str, Any]], group_tags: list[str]) -> None:
    request = item.get("request")
    name = str(item.get("name") or "").strip()
    next_tags = [*group_tags, name] if name and not request else group_tags

    if isinstance(request, dict):
        method = str(request.get("method") or "GET").upper()
        url_value = request.get("url")
        path, server_url = _extract_path_and_server_from_url_field(url_value)
        headers = _headers_from_pairs(request.get("header"))
        body = _postman_body_to_payload(request.get("body") or {})
        operations.append(
            _build_operation(
                method=method,
                path=path,
                summary=name or method,
                description=str(item.get("description") or ""),
                tags=group_tags[-1:] if group_tags else [],
                request_body=body,
                parameters=_params_from_query(url_value),
                servers=[{"url": server_url}] if server_url else [],
                headers=headers,
            )
        )

    for child in item.get("item", []) or []:
        if isinstance(child, dict):
            _walk_postman_items(child, operations, next_tags)


def _parse_har(data: dict[str, Any], *, source_name: str) -> dict[str, Any] | None:
    entries = (((data.get("log") or {}).get("entries")) or [])
    if not entries:
        return None

    operations: list[dict[str, Any]] = []
    seen: set[tuple[str, str]] = set()
    for entry in entries:
        request = entry.get("request") or {}
        url = request.get("url")
        if not url:
            continue
        parsed = urlparse(url)
        method = str(request.get("method") or "GET").upper()
        path = parsed.path or "/"
        key = (method, path)
        if key in seen:
            continue
        seen.add(key)
        query_params = []
        for item in request.get("queryString") or []:
            if isinstance(item, dict) and item.get("name"):
                query_params.append({
                    "name": item.get("name"),
                    "in": "query",
                    "required": False,
                    "schema": {"type": "string"},
                    "example": item.get("value", ""),
                })
        operations.append(
            _build_operation(
                method=method,
                path=path,
                summary=request.get("comment") or f"{method} {path}",
                parameters=query_params,
                responses={"200": {"description": "Captured from HAR"}},
                servers=[{"url": f"{parsed.scheme}://{parsed.netloc}"}] if parsed.scheme and parsed.netloc else [],
            )
        )

    if not operations:
        return None
    return _build_api_info_from_operations(operations, title=f"HAR Import {source_name}")


def _parse_generic_request_tree(data: dict[str, Any], *, source_name: str) -> dict[str, Any] | None:
    operations: list[dict[str, Any]] = []
    seen: set[tuple[str, str]] = set()

    def walk(node: Any, inherited_name: str = "") -> None:
        if isinstance(node, dict):
            request = node.get("request") if isinstance(node.get("request"), dict) else node
            method = str(request.get("method") or request.get("apiMethod") or "").upper()
            raw_url = request.get("url") or request.get("path") or request.get("apiPath")
            if method in HTTP_METHODS and raw_url:
                path, server_url = _extract_path_and_server_from_url_field(raw_url)
                key = (method, path)
                if key not in seen:
                    seen.add(key)
                    operations.append(
                        _build_operation(
                            method=method,
                            path=path,
                            summary=str(node.get("name") or request.get("name") or inherited_name or f"{method} {path}"),
                            description=str(node.get("description") or request.get("description") or ""),
                            tags=[inherited_name] if inherited_name else [],
                            servers=[{"url": server_url}] if server_url else [],
                        )
                    )
            next_name = str(node.get("name") or inherited_name or "")
            for value in node.values():
                walk(value, next_name)
        elif isinstance(node, list):
            for item in node:
                walk(item, inherited_name)

    walk(data)
    if not operations:
        return None
    title = str(data.get("name") or data.get("projectName") or data.get("title") or f"Imported APIs {source_name}")
    return _build_api_info_from_operations(operations, title=title)


async def _discover_openapi_from_html(base_url: str, html: str, client: httpx.AsyncClient) -> dict[str, Any] | None:
    candidates = _extract_html_doc_candidates(base_url, html)
    for candidate in candidates:
        try:
            response = await client.get(candidate)
            response.raise_for_status()
            openapi_info = _try_parse_openapi(response.content, suffix=Path(urlparse(str(response.url)).path).suffix.lower(), source_name=candidate)
            if openapi_info:
                return openapi_info
        except httpx.HTTPError:
            continue
    return None


def _extract_html_doc_candidates(base_url: str, html: str) -> list[str]:
    candidates: list[str] = []
    for path in DOC_DISCOVERY_PATHS:
        candidates.append(urljoin(base_url, path))

    if BeautifulSoup is not None:
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup.find_all(["a", "link", "script"]):
            value = tag.get("href") or tag.get("src")
            if value and any(token in value.lower() for token in ("openapi", "swagger", "api-docs")):
                candidates.append(urljoin(base_url, value))

    for matched in re.findall(r'["\']([^"\']+(?:openapi|swagger|api-docs)[^"\']*)["\']', html, flags=re.IGNORECASE):
        candidates.append(urljoin(base_url, matched))

    seen = set()
    result = []
    for item in candidates:
        if item not in seen:
            seen.add(item)
            result.append(item)
    return result


def _extract_text(content: bytes, *, suffix: str, source_name: str) -> str:
    if suffix == ".docx":
        return _parse_word(content)
    if suffix in {".xlsx", ".xls"}:
        return _parse_excel(content, suffix)
    if suffix == ".csv":
        return content.decode("utf-8", errors="ignore")
    if suffix in {".md", ".txt"}:
        return content.decode("utf-8", errors="ignore")
    if suffix in {".html", ".htm"}:
        return _html_to_text(content)
    return content.decode("utf-8", errors="ignore")


def _html_to_text(content: bytes) -> str:
    raw = content.decode("utf-8", errors="ignore")
    if BeautifulSoup is not None:
        soup = BeautifulSoup(raw, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
    else:
        text = re.sub(r"<script[^>]*>.*?</script>", "", raw, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<style[^>]*>.*?</style>", "", text, flags=re.DOTALL | re.IGNORECASE)
        text = re.sub(r"<[^>]+>", " ", text)
        text = html.unescape(text)
    return "\n".join([line.strip() for line in text.splitlines() if line.strip()])


def _parse_word(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise ValueError("当前环境缺少 python-docx，无法解析 Word 文档") from exc

    doc = Document(io.BytesIO(content))
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text for cell in row.cells)
            if row_text.strip():
                texts.append(row_text)
    return "\n".join(texts)


def _parse_excel(content: bytes, suffix: str) -> str:
    if suffix == ".xls":
        try:
            import xlrd
        except ImportError as exc:  # pragma: no cover
            raise ValueError("当前环境缺少 xlrd，无法解析 xls 文档") from exc

        wb = xlrd.open_workbook(file_contents=content)
        texts = []
        for sheet_name in wb.sheet_names():
            ws = wb.sheet_by_name(sheet_name)
            for row_idx in range(ws.nrows):
                row = ws.row_values(row_idx)
                row_text = " | ".join(str(cell) if cell else "" for cell in row)
                if row_text.strip():
                    texts.append(row_text)
        return "\n".join(texts)

    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover
        raise ValueError("当前环境缺少 openpyxl，无法解析 Excel 文档") from exc

    wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
    texts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        for row in ws.iter_rows(values_only=True):
            row_text = " | ".join(str(cell) if cell is not None else "" for cell in row)
            if row_text.strip():
                texts.append(row_text)
    wb.close()
    return "\n".join(texts)


def _build_api_info_from_text(text: str, *, title: str, server_url: str = "") -> dict[str, Any]:
    operations = _extract_operations_from_text(text, fallback_server=server_url)
    if not operations:
        raise ValueError("未能从内容中识别接口信息，请提供 OpenAPI、Postman、HAR 或包含 METHOD + PATH 的接口文档")
    return _build_api_info_from_operations(operations, title=title)


def _extract_operations_from_text(text: str, *, fallback_server: str = "") -> list[dict[str, Any]]:
    operations: list[dict[str, Any]] = []
    seen: set[tuple[str, str]] = set()
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    for line in lines:
        normalized = re.sub(r"\s+", " ", line)
        method_match = re.search(r"\b(GET|POST|PUT|PATCH|DELETE|OPTIONS|HEAD)\b\s+([^\s|]+)", normalized, flags=re.IGNORECASE)
        if method_match:
            method = method_match.group(1).upper()
            raw_target = method_match.group(2).strip("`'\"")
            path, server_url = _path_and_server_from_string(raw_target, fallback_server=fallback_server)
            summary = normalized[len(method_match.group(0)):].strip(" -|:\t")
            key = (method, path)
            if key not in seen:
                seen.add(key)
                operations.append(
                    _build_operation(
                        method=method,
                        path=path,
                        summary=summary or f"{method} {path}",
                        servers=[{"url": server_url}] if server_url else [],
                    )
                )
            continue

        if re.match(r"^(https?://\S+|/\S+)$", normalized):
            path, server_url = _path_and_server_from_string(normalized, fallback_server=fallback_server)
            key = ("GET", path)
            if key not in seen:
                seen.add(key)
                operations.append(
                    _build_operation(
                        method="GET",
                        path=path,
                        summary=f"GET {path}",
                        servers=[{"url": server_url}] if server_url else [],
                    )
                )

    return operations


def _build_api_info_from_operations(operations: list[dict[str, Any]], *, title: str) -> dict[str, Any]:
    path_map: dict[str, list[dict[str, Any]]] = {}
    tags: set[str] = set()
    servers: list[dict[str, Any]] = []
    seen_servers: set[str] = set()

    for operation in operations:
        path_map.setdefault(operation["path"], []).append(operation)
        for tag in operation.get("tags") or []:
            if tag:
                tags.add(tag)
        for server in operation.pop("_servers", []):
            url = str((server or {}).get("url") or "").strip()
            if url and url not in seen_servers:
                seen_servers.add(url)
                servers.append({"url": url})

    return {
        "info": {"title": title, "description": "", "version": ""},
        "servers": servers,
        "paths": [{"path": path, "operations": ops} for path, ops in sorted(path_map.items())],
        "components": {},
        "security": [],
        "tags": [{"name": tag} for tag in sorted(tags)],
    }


def _build_operation(
    *,
    method: str,
    path: str,
    summary: str,
    description: str = "",
    tags: list[str] | None = None,
    parameters: list[dict[str, Any]] | None = None,
    request_body: dict[str, Any] | None = None,
    responses: dict[str, Any] | None = None,
    security: list[dict[str, Any]] | None = None,
    servers: list[dict[str, Any]] | None = None,
    headers: dict[str, str] | None = None,
) -> dict[str, Any]:
    params = list(parameters or [])
    if headers:
        for key, value in headers.items():
            params.append({
                "name": key,
                "in": "header",
                "description": "",
                "required": False,
                "schema": {"type": "string"},
                "example": value,
            })

    return {
        "method": method.upper(),
        "path": path or "/",
        "operation_id": _operation_id(method, path),
        "summary": summary,
        "description": description,
        "tags": [tag for tag in (tags or []) if tag],
        "parameters": params,
        "request_body": request_body or {},
        "responses": responses or {},
        "security": security or [],
        "_servers": servers or [],
    }


def _ensure_api_info_shape(api_info: dict[str, Any]) -> dict[str, Any]:
    return {
        "info": api_info.get("info", {}),
        "servers": api_info.get("servers", []),
        "paths": api_info.get("paths", []),
        "components": api_info.get("components", {}),
        "security": api_info.get("security", []),
        "tags": api_info.get("tags", []),
    }


def _extract_path_and_server_from_url_field(url_value: Any) -> tuple[str, str]:
    if isinstance(url_value, str):
        return _path_and_server_from_string(url_value)

    if isinstance(url_value, dict):
        raw = url_value.get("raw")
        if isinstance(raw, str) and raw:
            return _path_and_server_from_string(raw)
        host = url_value.get("host") or []
        path = url_value.get("path") or []
        protocol = url_value.get("protocol") or "http"
        server = ""
        if host:
            server = f"{protocol}://{'.'.join([str(part) for part in host if part])}"
        normalized_path = "/" + "/".join([str(part).strip("/") for part in path if str(part).strip("/")]) if path else "/"
        return normalized_path, server

    return "/", ""


def _path_and_server_from_string(value: str, *, fallback_server: str = "") -> tuple[str, str]:
    if value.startswith("http://") or value.startswith("https://"):
        parsed = urlparse(value)
        return parsed.path or "/", f"{parsed.scheme}://{parsed.netloc}"
    if value.startswith("/"):
        return value, fallback_server
    return f"/{value.lstrip('/')}", fallback_server


def _server_from_url(url: str) -> str:
    parsed = urlparse(url)
    if parsed.scheme and parsed.netloc:
        return f"{parsed.scheme}://{parsed.netloc}"
    return ""


def _headers_from_pairs(items: Any) -> dict[str, str]:
    if not isinstance(items, list):
        return {}
    headers = {}
    for item in items:
        if isinstance(item, dict) and item.get("key"):
            headers[str(item["key"])] = str(item.get("value", ""))
    return headers


def _params_from_query(url_value: Any) -> list[dict[str, Any]]:
    query_params = []
    if isinstance(url_value, dict):
        for item in url_value.get("query") or []:
            if isinstance(item, dict) and item.get("key"):
                query_params.append({
                    "name": item.get("key"),
                    "in": "query",
                    "description": "",
                    "required": False,
                    "schema": {"type": "string"},
                    "example": item.get("value", ""),
                })
    return query_params


def _postman_body_to_payload(body: dict[str, Any]) -> dict[str, Any]:
    mode = body.get("mode")
    if mode == "raw":
        return {
            "description": "",
            "content": {
                "application/json": {
                    "example": body.get("raw", ""),
                }
            },
        }
    if mode == "urlencoded":
        return {
            "description": "",
            "content": {
                "application/x-www-form-urlencoded": {
                    "example": {item.get("key"): item.get("value") for item in body.get("urlencoded", []) if isinstance(item, dict) and item.get("key")},
                }
            },
        }
    return {}


def _operation_id(method: str, path: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "_", f"{method}_{path}").strip("_")
    return slug or method.lower()
